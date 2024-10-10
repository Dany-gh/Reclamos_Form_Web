# Primera Version: 240624
# Se ejecuta desde el terminal: pythonleer.py agua
# Se ejecuta desde el terminal: pythonleer.py luz 

# ---------- SE CONECTA USANDO oauth2 ---------------------------------------------------------------
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from google.oauth2 import service_account

import jwt
import time

from googleapiclient.errors import HttpError
# Esto importe cuando salio error con el json
from google.auth.exceptions import DefaultCredentialsError
#
from pathlib import Path

# Para limpiar la pantalla
import os
# Para word
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
# Para
import shutil
# Para leer datos por consola
import sys
# Para Depurar un programa
import pdb
# Para sacar fecha de hoy
from datetime import datetime
# Para usar otra manera de crear documentos de word. Sin usar Plantilla
from docx import Document
# Para enviar correo
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

####################################################################################################
# CONFIGURACION DE USUARIOS
# ------- PARA EL GOOGLE SHEET
# Alcances necesarios para acceder a la API de Google Sheets
SCOPES = ['https://www.googleapis.com/auth/spreadsheets']

# Escribe aquí el ID de tu documento:
# ID o URL de la hoja de cálculo y el Rango
SPREADSHEET_ID = '1FLWBfOe_ZKTMOviNBR35aC4CkHDGOwN2djyBd-rO0Js'

global SHEET_NAME
SHEET_NAME='' # Aqui voy a tener con que hoja trabajo. si es de LUZ o es de AGUA
SHEET_NAME_REC_LUZ = 'RtaFormRecLuz'
SHEET_NAME_REC_AGUA = 'RtaFormRecAgua'
SHEET_NAME_PRUEBA = 'ReclamosRes055-20'

# -------- PARA EL WORD
# Ruta al fichero Excel
#EXCEL_PATH= '.\Inputs\People_Data.xlsx'

# Ruta de Salida
ruta_archivo_actual = os.path.abspath(__file__)
#a = os.path.dirname(ruta_archivo_actual) # C:\.....\Reclamos_form_web\src
#b = os.path.dirname(a) # C:\.....\Reclamo_form_web
carpeta_reclamos = os.path.dirname(os.path.dirname(ruta_archivo_actual))
# Ruta al archivo en el mismo directorio donde se está ejecutando el script
OUTPUT_PATH = os.path.join(carpeta_reclamos,'Outputs' ) # Apunta al directorio Outputs

# Ruta de Imagenes
IMAGE_PATH = './Inputs/Images'

# Ruta plantillas o Templates. Ficheros word
global WORD_TEMPLATE # Defino la plantilla general
WORD_TEMPLATE = ''
ES_WORD_TPL_PATH = r'.\Inputs\Templates\WordTemplate_ES.docx'
EN_WORD_TPL_PATH = r'.\Inputs\Templates\WordTemplate_EN.docx'
WORD_TPL_PRUEBA1 = '.\\Inputs\\Templates\\WordTemplate_Prueba1.docx'
WORD_TPL_PRUEBA_L2 = '.\\Inputs\\Templates\\TemplateRECLAMOS_LUZ2.docx'
WORD_TPL_PRUEBA_A2 = './Inputs/Templates/TemplateRECLAMOS_AGUA2.docx'

global nombre_archivo
nombre_archivo = ''
global tipo_Reclamo
tipo_Reclamo = ''
global ultimo_color_usado # Pra saber que color se uso en la ultima pintada de celdas
ultimo_color_usado = ''

# Definimos los códigos de colores ANSI
class TextColor:
    RED = '\033[31m'
    GREEN = '\033[32m'
    YELLOW = '\033[33m'
    BLUE = '\033[34m'
    MAGENTA = '\033[35m'
    CYAN = '\033[36m'
    RESET = '\033[0m'  # Resetea el color al predeterminado
# Ejemplo de uso
#print(f"{TextColor.RED}Este texto es rojo.{TextColor.RESET}")

#==============================================================================================================================
# Limpia pantalla
def clear_screen():
    # Detecta el sistema operativo
    if os.name == 'nt':  # Para Windows
        os.system('cls')
    else:  # Para Unix/Linux/MacOS
        os.system('clear')
#------------------------------------------------------------------------------------------------------------------------------

#==============================================================================================================================
# Defino si es un Reclamo de LUZ o de AGUA.
def TipoReclamo(reclamo):
    global SHEET_NAME
    global WORD_TEMPLATE
    print(f"Tipo de Reclamo: {TextColor.GREEN}{reclamo}{TextColor.RESET}")
    if reclamo == 'LUZ':
        # Es un reclamo de LUZ
        SHEET_NAME = SHEET_NAME_REC_LUZ # Digo que trabajo con la Hoja de Luz
        WORD_TEMPLATE = WORD_TPL_PRUEBA_L2 # Uso esta Plantilla
        # Falta definir la planilla
    else:
        # Es un reclamo de AGUA
        SHEET_NAME = SHEET_NAME_REC_AGUA
        WORD_TEMPLATE = WORD_TPL_PRUEBA_A2
        #SHEET_NAME=SHEET_NAME_PRUEBA
        #exit(0)
    print(f"HOJA SELECCIONADA: {TextColor.GREEN}{SHEET_NAME}{TextColor.RESET}")
    print(f"PLANTILLA USADA: {TextColor.GREEN}{WORD_TEMPLATE}{TextColor.RESET}")
#------------------------------------------------------------------------------------------------------------------------------

#==============================================================================================================================
# Encuentra la PRIMERA FILA no leída (de la primera columna, NO es VERDE o NO es AMARILLA)
def find_first_unread_row(rows):
    global ultimo_color_usado

    for i, row in enumerate(rows):
        cell = row['values'][0]
        if 'effectiveFormat' in cell:
            background = cell['effectiveFormat']['backgroundColor']
            # Verifica si el color de la celda es VERDE (0,1,0) o AMARILLA (1,1,0)
            if not (background.get('red', 0) == 0 and background.get('green', 0) == 1 and background.get('blue', 0) == 0) and not (background.get('red', 0) == 1 and background.get('green', 0) == 1 and background.get('blue', 0) == 0):
                # No es VERDE o NO es AMARILLA la celda
                print(f"{TextColor.BLUE} Primera Fila Sin Leer: {TextColor.RESET}",i+2)
                return i + 2
            else:
                if (background.get('red', 0) == 0 and background.get('green', 0) == 1 and background.get('blue', 0) == 0):
                    ultimo_color_usado = 'VERDE'
                elif (background.get('red', 0) == 1 and background.get('green', 0) == 1 and background.get('blue', 0) == 0):
                    ultimo_color_usado = 'AMARILLO'
        else:
            #
            #return i + 2
            return None
                 
    #print("\033[34m Primera Fila:\033[0m",i+2)
    return None
#------------------------------------------------------------------------------------------------------------------------------

#==============================================================================================================================
# Devuelve la CANTIDAD de FILAS NO leídas (de la primera columna NO es VERDE y NO es AMARILLA)
def find_cant_unread_row(rows):
    cont_Filas_No_Verdes_No_Amarillas = 0
    for i, row in enumerate(rows):
        cell = row['values'][0]
        if 'effectiveFormat' in cell:
            background = cell['effectiveFormat']['backgroundColor']
            # Verifica si el color es: VERDE o AMARILLO
            if not (background.get('red', 0) == 0 and background.get('green', 0) == 1 and background.get('blue', 0) == 0) and not (background.get('red', 0) == 1 and background.get('green', 0) == 1 and background.get('blue', 0) == 0):
                # No es VERDE y no es AMARILLA la celda
                cont_Filas_No_Verdes_No_Amarillas +=1
                '''
                if (background.get('red', 0) == 1 and background.get('green', 0) == 0 and background.get('blue', 0) == 0):
                    # Si la celda es de COLOR ROJO sale.
                    return cont_Filas_No_Verdes
                '''
        else:
            #
            return cont_Filas_No_Verdes_No_Amarillas
    
    #print(f"{TextColor.BLUE}Cant. Filas No Verdes\\Amarillas: {TextColor.RESET} {cont_Filas_No_Verdes_No_Amarillas}")
    print(f"{TextColor.BLUE}Cant. Filas No Verdes\\Amarillas: {TextColor.RESET}", cont_Filas_No_Verdes_No_Amarillas)
    return cont_Filas_No_Verdes_No_Amarillas
#------------------------------------------------------------------------------------------------------------------------------

#==============================================================================================================================
# Imprime los valores de cada fila
def print_row_data(row):
    """
    Imprime los datos de una fila.
    :param row: Lista con los datos de la fila.
    """
    print('-' * 80)  # Línea separadora entre filas
    print(', '.join(row))
    print('-' * 80)  # Línea separadora entre filas
#------------------------------------------------------------------------------------------------------------------------------

#==============================================================================================================================
# Rutina para eliminar y crear carpeta
def EliminarCrearCarpetas(path):
    #Verificar si la carpeta existe y elimninarla
    if(os.path.exists(path)):
        shutil.rmtree(path)
        
    # Crear carpeta
    os.mkdir(OUTPUT_PATH)
#------------------------------------------------------------------------------------------------------------------------------

#==============================================================================================================================
# Rutina para crear un fichero word para cada persona 
# ESTA FORMA SI USA PLANTILLA PRE DEFINIDA. 
def CrearWordPersonas(df_pers):
    # Iteramos sobre cada Persona
    for r_idx, r_val in enumerate(df_pers):
        # Cargar plantilla
        l_tpl=WORD_TEMPLATE   # Plantilla o Template que se va a usar.
        '''
        if (r_val['Idioma'] == 'ES'):
            l_tpl=ES_WORD_TPL_PATH
        elif (r_val['Idioma'] == 'EN'):
            l_tpl=EN_WORD_TPL_PATH
        '''
        # Procesamos la plantilla
        docx_tpl=DocxTemplate(l_tpl)

        # Añadir imagen grafico circular y de barra
        #img_path = IMAGE_PATH + '\\' + r_val['Imagen']
        #img = InlineImage(docx_tpl, img_path, height=Mm(15))

        # Crear contexto
        # word : Google Sheet
        
        context = {
            'name': r_val['Nombre'],
            'surname1': r_val['Apellido'],
            'surname2': r_val['Telefono de Contacto'], # En este me da error
            'edad': r_val['Correo Electrónico'],
            #'picture': img,
        }
        # Crear el contexto (ChatGPT)
        #contexto = {'items': [{'indice': i, 'valor': v} for i, v in enumerate(df_pers)]}

        # Verificar el contexto
        print("Contexto:", context)

        # Renderizamos usando el contexto creado
        docx_tpl.render(context)

        # Guardamos el documento
        nombre_doc = 'Documento_' + r_val['Apellido'].upper() + '_' + r_val['Nombre'] + '.docx'
        '''
        if(pd.notna(r_val['Apellido2'])):
            nombre_doc = 'Documento_' + r_val['Apellido1'].upper() + '_' + r_val['Apellido2'].upper() + '_' + r_val['Nombre'] + '.docx'
        else:
            nombre_doc = 'Documento_' + r_val['Apellido1'].upper() + '_' + r_val['Nombre'] + '.docx'
        docx_tpl.save(OUTPUT_PATH + '\\' + nombre_doc)
        '''
        docx_tpl.save(OUTPUT_PATH + '\\' + nombre_doc)
#------------------------------------------------------------------------------------------------------------------------------

#==============================================================================================================================
# Rutina para crear un fichero word para TODAS las personas 
# ESTA FORMA SI USA PLANTILLA PRE DEFINIDA. 
def crea_documento_unico(datos_para_diccionario):
    try:
        # Convertir la lista a una lista de diccionarios
        #lista_diccionarios = [{"nombre": item[0], "edad": item[1]} for item in datos]
        
        # Cargar la plantilla
        doc = DocxTemplate(WORD_TEMPLATE)
        
        # Crear el contexto con todos los datos
        contexto = {'items': datos_para_diccionario}
        
        # Verificar el contexto
        print("Contexto:", contexto)
        # Renderizar el documento con el contexto
        doc.render(contexto)
        
        # Obtener la fecha actual
        fecha_actual = datetime.now()
        # Extraer día, mes y año
        dia = fecha_actual.day
        mes = fecha_actual.month
        anio = fecha_actual.year
        
        # Formatear el nombre del archivo
        nombre_archivo = f"RECLAMO_{anio}{mes:02d}{dia:02d}.docx"

        if SHEET_NAME == SHEET_NAME_REC_LUZ:
            # Formatear el nombre del archivo
            nombre_doc = f"RECLAMOS LUZ_{anio}{mes:02d}{dia:02d}.docx"

        elif SHEET_NAME == SHEET_NAME_REC_AGUA:
            # Formatear el nombre del archivo
            nombre_doc = f"RECLAMOS AGUA_{anio}{mes:02d}{dia:02d}.docx"

        # Guardar el documento
        doc.save(OUTPUT_PATH + '\\' + nombre_doc)
        
        print(f"Documento guardado como: {nombre_doc}")
    
    except Exception as e:
        print(f"{TextColor.RED}Ocurrió un error: {TextColor.RESET}", e)
#------------------------------------------------------------------------------------------------------------------------------

#==============================================================================================================================
# Rutina para crear un fichero word para TODAS las personas (OTRA MANERA)
# ESTA FORMA NO USA PLANTILLA PRE DEFINIDA. 
def OtraFormaCrearWord(datos_para_diccionario):
    global nombre_archivo
    
    # Obtener la fecha actual
    fecha_actual = datetime.now()
    anio = fecha_actual.year
    mes = fecha_actual.month
    dia = fecha_actual.day
    # Obtener los dos últimos dígitos del año
    anio = anio % 100

    # Obtener la hora actual del sistema
    hora_actual = datetime.now()
    # Extraer hora, minutos y segundos
    hora = hora_actual.hour
    minutos = hora_actual.minute

    # Formatear el nombre del archivo
    if SHEET_NAME == SHEET_NAME_REC_LUZ:
        nombre_archivo = f"RECLAMOS LUZ_{anio:02d}{mes:02d}{dia:02d}_{hora:02d}{minutos:02d}.docx"
    
        # Crear un documento de Word
        documento = Document()

        # Definir contador de hojas
        contador_hojas = 1

        # Iterar sobre cada diccionario en la lista y agregarlo al documento
        for indice, diccionario in enumerate(datos_para_diccionario):
            # Agregar el contenido del diccionario al documento
            documento.add_paragraph(f"RECLAMO DE LUZ NRO: {contador_hojas}")
            documento.add_paragraph(f"MARCA: {diccionario['Marca_Temporal']}")
            documento.add_paragraph(f"Apellido: {diccionario['Apellido']}")
            documento.add_paragraph(f"NOMBRE: {diccionario['Nombre']}")
            documento.add_paragraph(f"DNI: {diccionario['DNI']}")
            documento.add_paragraph(f"NRO. TEL: {diccionario['Nro_de_Telefono']}")
            documento.add_paragraph(f"CORREO: {diccionario['E_Mail']}")
            documento.add_paragraph(f"DOMICILIO: {diccionario['Domicilio']}")
            documento.add_paragraph(f"NRO. SUMINISTRO: {diccionario['Nro_de_Suministro']}")
            documento.add_paragraph(f"Tipo de Reclamo: {diccionario['Tipo_de_Reclamo']}")
            documento.add_paragraph(f"DESCRIPCION: {diccionario['Descripcion_Reclamo']}")

            # Agregar un salto de página después de cada elemento
            documento.add_page_break()

            # Incrementar el contador de hojas
            contador_hojas += 1
    
    elif SHEET_NAME == SHEET_NAME_REC_AGUA:
        nombre_archivo = f"RECLAMOS AGUA_{anio:02d}{mes:02d}{dia:02d}_{hora:02d}{minutos:02d}.docx"
        #nombre_archivo = f"RECLAMOS AGUA_{anio}{mes:02d}{dia:02d}.docx"

        # Crear un documento de Word
        documento = Document()

        # Definir contador de hojas
        contador_hojas = 1

        # Iterar sobre cada diccionario en la lista y agregarlo al documento
        for indice, diccionario in enumerate(datos_para_diccionario):
            # Agregar el contenido del diccionario al documento
            documento.add_paragraph(f"RECLAMO DE AGUA NRO: {contador_hojas}")
            documento.add_paragraph(f"MARCA: {diccionario['Marca_Temporal']}")
            documento.add_paragraph(f"Apellido: {diccionario['Apellido']}")
            documento.add_paragraph(f"NOMBRE: {diccionario['Nombre']}")
            documento.add_paragraph(f"DNI: {diccionario['DNI']}")
            documento.add_paragraph(f"NRO. TEL: {diccionario['Nro_de_Telefono']}")
            documento.add_paragraph(f"CORREO: {diccionario['E_Mail']}")
            documento.add_paragraph(f"DOMICILIO: {diccionario['Domicilio']}")
            documento.add_paragraph(f"NRO. SUMINISTRO: {diccionario['Nro_de_Suministro']}")
            documento.add_paragraph(f"DESCRIPCION: {diccionario['Descripcion_Reclamo']}")

            # Agregar un salto de página después de cada elemento
            documento.add_page_break()

            # Incrementar el contador de hojas
            contador_hojas += 1
    
    # Guardar el documento
    documento.save(OUTPUT_PATH + '\\' + nombre_archivo)

    print(f"Archivo {TextColor.GREEN}'{nombre_archivo}'{TextColor.RESET} creado exitosamente.")    
#------------------------------------------------------------------------------------------------------------------------------

#==============================================================================================================================
# Rutina: para enviar correo
# Envia Correo con elemento adjunto
def Enviar_Correo(destinatario, asunto, cuerpo, archivo_adjunto, remitente, password):
    
    # Configuración de los datos de acceso a Gmail
    smtp_servidor = 'smtp.gmail.com'
    smtp_port = 587 # Puerto seguro TLS

    # Obtener directorio actual del script
    directorio_script = os.path.dirname(os.path.abspath(__file__))
    
    # Subir un nivel hacia la carpeta que contiene 'Outputs'
    directorio_padre = os.path.dirname(directorio_script)   
    
    # Construir ruta completa al archivo
    ruta_archivo = os.path.join(directorio_padre, 'Outputs', archivo_adjunto)
    
    # Configuración del mensaje
    mensaje = MIMEMultipart()
    mensaje['From'] = remitente
    mensaje['To'] = destinatario
    mensaje['Subject'] = asunto

    # Adjuntar cuerpo del mensaje
    # Agregar el cuerpo del correo al mensaje
    mensaje.attach(MIMEText(cuerpo, 'plain'))

    # Adjuntar archivo
    with open(ruta_archivo, 'rb') as adjunto:
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(adjunto.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename= {archivo_adjunto}')
        mensaje.attach(part)
    
    try:
        # Establecer conexión con el servidor SMTP
        servidor_smtp = smtplib.SMTP(host=smtp_servidor, port=smtp_port)
        servidor_smtp.starttls() # Habilitar seguridad TLS

        # Autenticación
        servidor_smtp.login(remitente, password)

        # Enviar el correo a cada destinatario
        text = mensaje.as_string()
        servidor_smtp.sendmail(remitente, destinatario, text)
        print(f'{TextColor.BLUE}Correo Enviado Con Exito a: {destinatario}{TextColor.RESET}') # Imprime en color azul

        # Envío del correo
        #servidor_smtp.sendmail(remitente, destinatario, mensaje.as_string())
    except Exception as e:
        print(f'{TextColor.RED}Error al enviar correo: {TextColor.RESET} {str(e)}')
    finally:
        # Cerrar conexión
        servidor_smtp.quit()
#-----------------------------------------------------------------------------------------------------------------------------

# ==============================================================================================================================
# Rutina: para chequear el token. 
def chequear_token(credencial):
    # Obtener el token JWT
    jwt_token = credencial._make_authorization_grant_assertion()
    # Decodificar el token JWT para ver los campos 'iat' y 'exp'
    decoded_token = jwt.decode(jwt_token, options={"verify_signature": False})
    # Imprimir los valores de 'iat' y 'exp'
    print(f"Issued at (iat): {time.strftime('%Y-%m-%d %H:%M:%S', time.gmtime(decoded_token['iat']))}")
    print(f"Expiration (exp): {time.strftime('%Y-%m-%d %H:%M:%S', time.gmtime(decoded_token['exp']))}")


# ###########################################################################################################################
# RUTINA PRINCIPAL
def main():
    global SHEET_NAME
    global nombre_archivo # Nombre del archivo word generado, donde estan los reclamos.
    global ultimo_color_usado # Ultimo color que use para pintar celdas.

    # Ruta al archivo de credenciales .JSON
    current_dir = Path(__file__).parent
    # Cargo archivo .json
    KEY=current_dir/'clave_Reclamos_Form_Web.json'
    try:
        # Autenticación y acceso a la hoja de cálculo
        creds = None
        creds = service_account.Credentials.from_service_account_file(KEY, scopes=SCOPES)
        #chequear_token(creds)
        '''
        # ====================================================================================================
        # Obtener el token JWT
        jwt_token = creds._make_authorization_grant_assertion()
        # Decodificar el token JWT para ver los campos 'iat' y 'exp'
        decoded_token = jwt.decode(jwt_token, options={"verify_signature": False})
        # Imprimir los valores de 'iat' y 'exp'
        print(f"Issued at (iat): {time.strftime('%Y-%m-%d %H:%M:%S', time.gmtime(decoded_token['iat']))}")
        print(f"Expiration (exp): {time.strftime('%Y-%m-%d %H:%M:%S', time.gmtime(decoded_token['exp']))}")
        # /////////////////////////////////////////////////////////////////////////////////////////////////////
        '''
        service = build('sheets', 'v4', credentials=creds)

        # Llamada a la API
        sheet = service.spreadsheets()
        #pdb.set_trace()
        
        '''
        # Obtén el rango de datos existente en la hoja, por medio de la columna A
        result = sheet.values().get(spreadsheetId=SPREADSHEET_ID, range=f'{SHEET_NAME}!A:A').execute()
        '''

        '''
        result=
        range: NombreHoja!A1:A139
        majorDimension:ROWS
        values: [[Titulo Celda de A1],[],[],[],.......[valor de la celda A139 en este caso]] (es una lista de lista)
        '''

        '''
        print(f"{TextColor.BLUE}Resul: {TextColor.RESET}") # Imprime en color azul
        print(result) # Imprime el diccionario
                
        rows = result.get('values', []) # Aqui solo son las filas (rows) de la primera columna (A). Hasta la ultima fila que tiene valor
        print(f"{TextColor.BLUE} {len(rows)} : Filas (rows) Recuperadas.{TextColor.RESET}")
        # Saco la cantidad de filas que tienen datos, per no sabemos si estan pintadas o no.
        num_rows = len(result.get('values', [])) - 1 # Descuento la cabecera
        '''

        # Otra Manera de Sacar el Rango donde hay DATOS.
        result = sheet.values().get(spreadsheetId=SPREADSHEET_ID, range=f'{SHEET_NAME}').execute()
        # Obtengo solo los Datos, incluido los nombres de columnas del encabezado.
        values = result.get('values', [])
        # Determina el rango donde estan los datos con valores.
        # Filas y Columnas.
        if values:
            num_rows = len(values) # Cantidad de Filas (incluido el encabezado)
            num_cols = len(values[0]) # Cantidad de Columnas, de la fila 0
            start_cell = f"A1" # Inicio del rango de datos
            end_cell = f"{chr(64 + num_cols)}{num_rows}" # Formato 'ColumnaNroFila'
            print(f"Los datos en: {TextColor.YELLOW}'{SHEET_NAME}'{TextColor.RESET} van desde: {TextColor.YELLOW}{start_cell}{TextColor.RESET} hasta: {TextColor.YELLOW}{end_cell}{TextColor.RESET}")
            num_rows = num_rows - 1 # Le saco el encabezado
        else:
            print(f"{TextColor.RED}No se encontraron datos en: '{SHEET_NAME}'{TextColor.RESET}")
            # Tendria que terminar el programa. COMO HAGO?

        # EL PROGRAMA CONTINUA, PERO NO SE SI SOLO TENGO DATOS DE ENCABEZADO O MAS FILAS.
        '''
        #--chatGPT-------------------------------------------------------------------------------
        # Obtiene los detalles de la hoja de cálculo
        spreadsheet = sheet.get(spreadsheetId=SPREADSHEET_ID).execute()

        # Busca el sheetId correspondiente al SHEET_NAME
        sheet_id = None
        for sheet in spreadsheet['sheets']:
            if sheet['properties']['title'] == SHEET_NAME:
                sheet_id = sheet['properties']['sheetId']
                break
        #---------------------------------------------------------------------------------
        '''

        '''
        #---------------------------------------------------------------------------------
        # Imprime los títulos de las hojas y sus IDs
        for sheet in spreadsheet['sheets']:
            title = sheet['properties']['title']
            sheet_id = sheet['properties']['sheetId']
            print(f"Sheet title: {title}, Sheet ID: {sheet_id}")
        #---------------------------------------------------------------------------------
        '''

        if num_rows == 0:
            # SOLO TENIA EL DATO DE ENCABEZADO.
            print(f"{TextColor.BLUE}No hay datos en la hoja: {SHEET_NAME}{TextColor.RESET}")
            # AQUI TENDRIA QUE TERMINAR EL PROGRAMA.
        else:
            # SI HAY DATOS.
            # Rango sin encabezado
            RANGE = f'{SHEET_NAME}!A2:A{num_rows+1}'  # Rango basado en el número de filas con datos, sin tener en cuenta si estan pintadas o no.
           
            #===========================================
            # Obtén los datos y el formato de la hoja
            #-------------------------------------------------------------------------------------------------------------------------------------------------
            # 1 Forma: El RANGE es sin el encabezado.
            result = sheet.get(spreadsheetId=SPREADSHEET_ID, ranges=RANGE, fields='sheets(data.rowData.values.effectiveFormat)').execute() # Con chatGPT. OK
            rows = result['sheets'][0]['data'][0]['rowData'] # Con chatGPT. OK
            # rowDat:{values:[{..}]}
            #-------------------------------------------------------------------------------------------------------------------------------------------------
            
            '''
            #-------------------------------------------------------------------------------------------------------------------------------------------------
            # De esta manera no especifico con que hoja (sheet) quiero trabajar. Por defecto toma la hoja 1.
            # 2 Forma:
            result = sheet.get(spreadsheetId=SPREADSHEET_ID, fields='sheets(data.rowData.values.effectiveFormat)').execute() # Con chatBlackbox. Me toma la sheet1. OK
            print(f"{TextColor.BLUE} Result: {TextColor.RESET}")
            print(result)
            rows = result.get('sheets')[0].get('data')[0].get('rowData') # Con chatblackbox. OK
            print(f"{TextColor.BLUE} rows: {TextColor.RESET}")
            print(rows)
            #-------------------------------------------------------------------------------------------------------------------------------------------------
            '''

            '''
            #-------------------------------------------------------------------------------------------------------------------------------------------------
            # 
            3 Forma:
            result = service.spreadsheets().values().batchGet(spreadsheetId=SPREADSHEET_ID, ranges=RANGE).execute() # De chatBlackbox. OK
            rows = result['sheets'][0]['data'][0]['rowData'] # No OK
            #-------------------------------------------------------------------------------------------------------------------------------------------------
            '''
            
            # Busco la primera fila que no esta pintada de VERDE o AMARILLA, del rango (RANGE) especificado.
            first_unread_row = find_first_unread_row(rows)
            if not (first_unread_row == None):
                cant_unread_row = find_cant_unread_row(rows) # Cantidad de filas no leidas (NO VERDE o NO AMARILLAS)
                last_unread_row = (cant_unread_row + first_unread_row) - 1 # Saco la ultima fila del Rango que NO tiene VERDE o AMARILLO. DONDE USO ESTE VALOR???
            else:
                # NO HAY REGISTROS NUEVOS PARA INFORMAR.
                first_unread_row = 0 # Le pongo valor cero para decir que no hay filas nuevas.
            
            if first_unread_row:
                # Procesa el registro en la primera fila no leída
                print(f"{TextColor.BLUE} Procesando desde la fila: {TextColor.RESET}", first_unread_row)
                
                # Este rango es donde esta mi informacion, nueva.
                if(SHEET_NAME == SHEET_NAME_REC_LUZ ):               
                    range_to_read = f'{SHEET_NAME}!A{first_unread_row}:{end_cell}'  # Ajusta el rango según sea necesario
                elif(SHEET_NAME == SHEET_NAME_REC_AGUA):
                    range_to_read = f'{SHEET_NAME}!A{first_unread_row}:{end_cell}'  # Ajusta el rango según sea necesario
                
                # ----CHATGPT
                rango_base=range_to_read.split('!')[1]  # Esto te dará 'CELDAx:CELDAy' CELDA=Letra, x e y numeros
                # Extraer el número de fila base
                fila_base = int(rango_base.split(':')[0][1:])  # Esto te dará nro de fila 
                # ---------------------------------------------------------------

                record = sheet.values().get(spreadsheetId=SPREADSHEET_ID, range=range_to_read).execute()
                print(f"{TextColor.BLUE}RECORD: {TextColor.RESET}") # Color AZUL
                # Obtengo la LISTA.
                datos_lista = record.get('values',[])
                
                # La primera fila contiene los nombres de las columnas
                columns = datos_lista[0][1] # apunta a la segunda columna en este caso 'Apellido' [Fila][Columna]
                valor = datos_lista[1:] # Me devuelve todos los datos a partir de la fila 1.
                
                '''
                # --- Chat GPT -----------------------------------------------
                # Remueve el carácter de nueva línea del quinto campo
                campo_corregido = datos_lista[8].replace("\n", "")

                # Crea una nueva lista con el campo corregido
                lista_corregida = datos_lista[:8] + [campo_corregido]

                # Imprime la nueva lista corregida
                print(lista_corregida)
                #-------------------------------------------------------------
                '''

                # --- Chat GPT -----------------------------------------------
                # Lo que quiero hacer es, eliminar los saltos de lineas en las celdas
                # Especifica el índice del campo que deseas corregir (en este caso, el campo 5 tiene índice 4)
                if SHEET_NAME == SHEET_NAME_REC_LUZ:
                    indice_campo_a_corregir =  9 # Descripcion Luz
                elif SHEET_NAME == SHEET_NAME_REC_AGUA:
                    indice_campo_a_corregir = 8 # Descrpcion Agua
                
                # Inicializa una nueva lista para almacenar las listas corregidas
                lista_corregida = []
                # Imprime el tipo de 
                #print(type(lista_corregida))

                # Recorre cada sublista en la lista de listas
                for sublista in datos_lista:
                    # Verifica si el campo a corregir es una cadena antes de intentar usar replace
                    if isinstance(sublista[indice_campo_a_corregir], str):
                        # Corrige el campo eliminando el carácter de nueva línea
                        sublista[indice_campo_a_corregir] = sublista[indice_campo_a_corregir].replace("\n", "")
                    
                    # Añade la sublista corregida a la nueva lista
                    lista_corregida.append(sublista)

                # Imprime la nueva lista de listas corregida
                #print(lista_corregida)

                datos_lista = lista_corregida
                # Imprime el tipo de 
                #print(type(datos_lista))

                #--------------------------------------------------------------- 
                if not datos_lista:
                    print(f'{TextColor.RED}No data found.{TextColor.RESET}')
                else:
                    # Itera sobre las filas y las imprime
                    for row in datos_lista:
                        print_row_data(row) # Llamo a la funcion print_row_data()
                #----------------------------------------------------------------    
                #pdb.set_trace() # Para depurar.
                if(SHEET_NAME == SHEET_NAME_REC_LUZ ):               
                    # Convertir la lista en una lista de diccionarios
                    datos_lista_diccionario = [{'Marca_Temporal': item[0], 
                                        'Apellido': item[1], 
                                        'Nombre' : item[2], 
                                        'DNI' : item[3],
                                        'Nro_de_Telefono' : item[4],
                                        'E_Mail' : item[5],
                                        'Domicilio' : item[6],
                                        'Nro_de_Suministro' : item[7],
                                        'Tipo_de_Reclamo' : item[8],
                                        'Descripcion_Reclamo' : item[9]} for item in datos_lista]
                elif(SHEET_NAME == SHEET_NAME_REC_AGUA):
                    # Convertir la lista en una lista de diccionarios
                    datos_lista_diccionario = [{'Marca_Temporal': item[0], 
                                        'Apellido': item[1], 
                                        'Nombre' : item[2], 
                                        'DNI' : item[3], 
                                        'Nro_de_Telefono' : item[4],
                                        'E_Mail' : item[5],
                                        'Domicilio' : item[6],
                                        'Nro_de_Suministro' : item[7],
                                        'Descripcion_Reclamo' : item[8]} for item in datos_lista]

                # Imprime el tipo de datos_diccionario
                #print(type(datos_lista_diccionario))
                
                '''
                #--- ChatGPT -----------------------------------------------------------------------
                # Recorre la lista de diccionario e imprime solo el campo especifico del diccionario
                # Define el campo específico que quieres imprimir
                campo_especifico = "Descripcion_Reclamo"
                # Verificar si datos_diccionario es realmente un diccionario
                
                if isinstance(datos_lista_diccionario, list):
                    for diccionario in datos_lista_diccionario:
                        # Verifica si el campo específico existe en el diccionario
                        if campo_especifico in diccionario:
                            print("")
                            print(diccionario[campo_especifico])
                        else:
                            print(f"{campo_especifico} no encontrado en el diccionario: {diccionario}")                    
                else:
                    print("datos_lista_diccionario, no es una LISTA ")
                #---------------------------------------------------------------------------------------
                '''

                '''
                # Convierte los datos a un DataFrame de Pandas
                if not datos_lista:
                    print(f'{TextColor.BLUE}No data found.{TextColor.RESET}')
                else:
                    # Asume que la primera fila de values contiene los nombres de las columnas
                    #df = pd.DataFrame(datos[1:], columns=datos[0])
                    # Muestra el DataFrame
                    #print(df)
                '''
                
                EliminarCrearCarpetas(OUTPUT_PATH)
                
                # Diferentes maneras de crear el WORD
                #CrearWordPersonas(datos_diccionario) # Crea una hoja de word por reclamo.
                #crea_documento_unico(datos_lista_diccionario) # Crea una hoja de word por multiples reclamos.
                OtraFormaCrearWord(datos_lista_diccionario) # Crea una hoja de word por multiples reclamos, sin plantilla.
                
                #----------Datos Para enviar Correo --------------------------------------------------
                # Obtener la fecha actual
                fecha_actual = datetime.now()
                # Extraer día, mes y año
                dia = fecha_actual.day
                mes = fecha_actual.month
                anio = fecha_actual.year
                
                archivo_adjunto =nombre_archivo
                remitente ='enrecat@catamarca.gov.ar'
                password ='enrecat16'

                if (tipo_Reclamo == 'LUZ'):
                    cuerpo ='Hola, adjunto te envio RECLAMOS DE ENERGIA al dia de la Fecha.'
                    asunto =f'RECLAMOS {tipo_Reclamo} AL DIA {dia:02d}-{mes:02d}-{anio}'
                    # Lista de destinatarios
                    destinatarios =['selememoises76@gmail.com','reclamosenergia@gmail.com']
                    #destinatarios =['daguirreie@yahoo.com.ar','daguirreie@gmail.com']
                    # Enviar el correo a cada destinatario individualmente
                    for destinatario in destinatarios:
                        Enviar_Correo(destinatario,asunto,cuerpo,archivo_adjunto,remitente,password)
                elif (tipo_Reclamo == 'AGUA'):
                    cuerpo ='Hola, adjunto te envio RECLAMOS DE AGUA al dia de la Fecha.'
                    asunto =f'RECLAMOS {tipo_Reclamo} AL DIA {dia:02d}-{mes:02d}-{anio}'
                    # Lista de destinatarios
                    destinatarios =['selememoises76@gmail.com','reclamosaguaycloacas@gmail.com']
                    #destinatarios =['daguirreie@yahoo.com.ar','daguirreie@gmail.com']
                    # Enviar el correo a cada destinatario individualmente
                    for destinatario in destinatarios:
                        Enviar_Correo(destinatario,asunto,cuerpo,archivo_adjunto,remitente,password)
                #--------------------------------------------------------------------------------------

                # oooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooo
                # PINTAR CELDAS
                
                # --- SACO sheet_id (chatGPT) ------------------------------------------------------------------------------
                # Obtiene los detalles de la hoja de cálculo. Saco el sheet_id
                spreadsheet = sheet.get(spreadsheetId=SPREADSHEET_ID).execute()
                # Busca el sheetId correspondiente al SHEET_NAME
                sheet_id = None
                for sheet in spreadsheet['sheets']:
                    if sheet['properties']['title'] == SHEET_NAME:
                        sheet_id = sheet['properties']['sheetId']
                        break
                #---------------------------------------------------------------------------------
                # Configuro valores para pintar las celdas, segun ultimo color usado
                if ultimo_color_usado == 'VERDE':
                    # Voy a usar AMARILLO
                    valorRed = 1
                    valorGreen = 1
                    valorBlue = 0
                elif ultimo_color_usado == 'AMARILLO':
                    # Voy a usar VERDE
                    valorRed = 0
                    valorGreen = 1
                    valorBlue = 0
                #--------------------------------------------------------------------------------- 
                # Marca la fila como leída, pintando la primera celda de VERDE O AMARILLA
                # SEGUN LA FILA ANTERIOR.
                # Tener en cuenta que la primer fila es la nro 0
                star_Fila = (first_unread_row - 1)
                end_Fila = (star_Fila + cant_unread_row)
                requests = [{
                    'updateCells': {
                        'range': {
                            'sheetId': sheet_id,    # Id de la hoja con la cual estamos trabajando
                            'startRowIndex': star_Fila, # Inicio del rango de filas.
                            'endRowIndex': end_Fila,   # Para abarcar mas filas (es exclusive)
                            'startColumnIndex': 0,
                            'endColumnIndex': 1,
                        },
                        'rows': [{
                            'values': [{
                                'userEnteredFormat': {
                                    'backgroundColor': {
                                        'red': valorRed,
                                        'green': valorGreen,
                                        'blue': valorBlue
                                    }
                                }
                            }]
                        }] * cant_unread_row, # Multiplica la cantidad de filas que quiero pintar
                        'fields': 'userEnteredFormat.backgroundColor'
                    }
                }]
                body = {'requests': requests}

                # Ejecuta la solicitud batchUpdate
                response = service.spreadsheets().batchUpdate(spreadsheetId=SPREADSHEET_ID, body=body).execute()
                #response = sheet.batchUpdate(spreadsheetId=SPREADSHEET_ID, body=body).execute() # Tengo error aqui
                print(f"\033[32m{cant_unread_row} :Filas Marcadas como leídas.{TextColor.RESET}")
                
                # oooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooooo
            else:
                print(f"{TextColor.BLUE}No hay registros nuevos para leer.{TextColor.RESET}")
    
    except DefaultCredentialsError as e:
        print(f"{TextColor.RED}Error en las credenciales (DefaultCredentialsError): {TextColor.RESET} {e}")
    
    except HttpError as error:
        print(f"{TextColor.RED}Un Error ha ocurrido (HttpError): {TextColor.RESET} {error}")
        return error
    
    except Exception as e:
        print(f"{TextColor.RED}Ocurrió un error (Exception): {TextColor.RESET}\n{e}")    

# /////////////////////////////////////////////////////////////////////////
if __name__ == '__main__':
    clear_screen()
    print(f"{TextColor.YELLOW}VERIFICACION DE RECLAMOS VIA WEB: DIRECCION AREA TECNICA. (En.Re.) - VERS.: 240909_241010{TextColor.RESET}")
    current_dir = Path(__file__).parent

    # Nombre del archivo que deseas verificar
    archivo = 'clave_Reclamos_Form_Web.json'

    # Ruta al archivo en el mismo directorio donde se está ejecutando el script
    ruta_al_archivo = os.path.join(current_dir, archivo)

    # Verifica si el archivo existe
    if os.path.exists(ruta_al_archivo):
        print(f"El archivo de credenciales: '{archivo}' {TextColor.GREEN}SI Existe.{TextColor.RESET}")
        # Verifica si se pasó algún argumento
        # sys.arg[0] = Tiene el Nombre del script
        #pdb.set_trace() # Para depurar el programa
        
        #'''
        if len(sys.argv) > 1:
            # Se paso algun argumento
            tipo_Reclamo = sys.argv[1].upper()
            print(f"Parametro Recibido: {TextColor.GREEN}{tipo_Reclamo}{TextColor.RESET}")
            
            # Aquí puedes tomar decisiones basadas en el valor de 'tipo_Reclamo' o parametro pasado.
            if tipo_Reclamo == "LUZ":
                print(f"Has seleccionado RECLAMO DE: {TextColor.GREEN}{tipo_Reclamo}{TextColor.RESET}")
                TipoReclamo(tipo_Reclamo)
                main()
            elif tipo_Reclamo == "AGUA":
                print(f"Has seleccionado RECLAMO DE: {TextColor.GREEN}{tipo_Reclamo}{TextColor.RESET}")
                TipoReclamo(tipo_Reclamo)
                main()
            else:
                print(f"{TextColor.RED}Opción no reconocida.{TextColor.RESET}")
        else:
            print(f"{TextColor.RED}No se proporcionó ningún Parametro. (NO SE SABE QUE TIPO DE RECLAMO ES [AGUA o LUZ]).{TextColor.RESET}")

        #'''
        
        '''
        # Es para probar el programa
        tipo_Reclamo='AGUA'
        TipoReclamo(tipo_Reclamo)
        main()
        '''
        print(f"{TextColor.MAGENTA}----- <<< FINAL PROGRAM >>> ---- {TextColor.RESET}")
        exit(0) # programa ha terminado correctamente, sin errores.
    else:
        print(f"El archivo de credenciales .JSON: '{archivo}' {TextColor.RED}NO Existe.{TextColor.RESET}")
        exit(0) # programa ha terminado correctamente, sin errores.
        